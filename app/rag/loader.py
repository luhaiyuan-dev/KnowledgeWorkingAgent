from pathlib import Path

from docx import Document as WordDocument
from langchain_core.documents import Document
from openpyxl import load_workbook
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".txt", ".md"}


class EnterpriseDocumentLoader:
    def __init__(self, allowed_root: str | Path) -> None:
        self.allowed_root = Path(allowed_root).resolve()

    def load_directory(self) -> list[Document]:
        documents: list[Document] = []
        for file_path in sorted(self.allowed_root.rglob("*")):
            if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                documents.extend(self.load_file(file_path))
        return documents

    def load_file(self, file_path: str | Path) -> list[Document]:
        safe_path = self._resolve_safe_path(file_path)
        extension = safe_path.suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式：{extension}")

        if extension in {".txt", ".md"}:
            return self._load_text(safe_path)
        if extension == ".pdf":
            return self._load_pdf(safe_path)
        if extension == ".docx":
            return self._load_word(safe_path)
        return self._load_excel(safe_path)

    def _resolve_safe_path(self, file_path: str | Path) -> Path:
        candidate = Path(file_path)
        if not candidate.is_absolute():
            candidate = self.allowed_root / candidate
        resolved_path = candidate.resolve()
        if resolved_path != self.allowed_root and self.allowed_root not in resolved_path.parents:
            raise PermissionError("文件路径超出知识库允许目录")
        if not resolved_path.is_file():
            raise FileNotFoundError(f"文件不存在：{resolved_path.name}")
        return resolved_path

    def _base_metadata(self, file_path: Path) -> dict[str, object]:
        return {
            "source": str(file_path),
            "source_name": file_path.name,
            "file_type": file_path.suffix.lower().lstrip("."),
            "access_scope": "internal",
        }

    def _load_text(self, file_path: Path) -> list[Document]:
        content = file_path.read_text(encoding="utf-8")
        return [Document(page_content=content, metadata=self._base_metadata(file_path))]

    def _load_pdf(self, file_path: Path) -> list[Document]:
        reader = PdfReader(str(file_path))
        documents: list[Document] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                metadata = self._base_metadata(file_path)
                metadata["page"] = page_number
                documents.append(Document(page_content=text, metadata=metadata))
        return documents

    def _load_word(self, file_path: Path) -> list[Document]:
        word_document = WordDocument(str(file_path))
        blocks: list[str] = []
        for paragraph in word_document.paragraphs:
            if paragraph.text.strip():
                blocks.append(paragraph.text.strip())
        for table_index, table in enumerate(word_document.tables, start=1):
            blocks.append(f"表格 {table_index}")
            for row in table.rows:
                cell_values = [cell.text.strip() for cell in row.cells]
                blocks.append(" | ".join(cell_values))
        content = "\n".join(blocks)
        return [Document(page_content=content, metadata=self._base_metadata(file_path))]

    def _load_excel(self, file_path: Path) -> list[Document]:
        workbook = load_workbook(filename=file_path, read_only=True, data_only=True)
        documents: list[Document] = []
        for worksheet in workbook.worksheets:
            rows = list(worksheet.iter_rows(values_only=True))
            if not rows:
                continue
            headers = [str(value or f"列{index + 1}") for index, value in enumerate(rows[0])]
            text_rows: list[str] = []
            for row_number, row in enumerate(rows[1:], start=2):
                fields: list[str] = []
                for header, value in zip(headers, row, strict=False):
                    if value is not None:
                        fields.append(f"{header}: {value}")
                if fields:
                    text_rows.append(f"第{row_number}行 | " + " | ".join(fields))
            metadata = self._base_metadata(file_path)
            metadata["sheet"] = worksheet.title
            documents.append(Document(page_content="\n".join(text_rows), metadata=metadata))
        workbook.close()
        return documents
